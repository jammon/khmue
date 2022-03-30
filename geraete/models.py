from datetime import date

from django.conf import settings
from django.contrib.auth.models import User
from django.db import models
from django.http import Http404
from docx import Document

# Create your models here.


class Company(models.Model):
    name = models.CharField("Name", max_length=50)

    class Meta:
        verbose_name = "Krankenhaus"
        verbose_name_plural = "Krankenhäuser"

    def __str__(self):
        return self.name


class DeviceCat(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="device_cats",
    )
    category = models.CharField("Kategorie", max_length=50)

    class Meta:
        verbose_name = "Gerätekategorie"
        verbose_name_plural = "Gerätekategorien"

    def __str__(self):
        return self.category


class Device(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="devices",
    )
    category = models.ForeignKey(
        DeviceCat, on_delete=models.CASCADE, verbose_name="Kategorie"
    )
    vendor = models.CharField("Hersteller", max_length=50)
    name = models.CharField("Name", max_length=50)

    class Meta:
        verbose_name = "Gerät"
        verbose_name_plural = "Geräte"

    def __str__(self):
        return f"{self.category}: {self.name} von {self.vendor} "

    @classmethod
    def get_for_instructor(cls, instructor_id, company_id):
        return Device.objects.filter(
            instructors__id=instructor_id, company__id=company_id
        )

    @classmethod
    def get_sorted(cls, request):
        return (
            company_s(request, cls)
            .select_related("category")
            .order_by("category__category", "vendor", "name")
        )


class ProfessionalGroup(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="professional_groups",
    )
    name = models.CharField("Name", max_length=50)
    devices = models.ManyToManyField(
        Device,
        verbose_name="Einzuweisende Geräte",
        blank=True,
        related_name="professional_groups",
    )

    class Meta:
        verbose_name = "Berufsgruppe"
        verbose_name_plural = "Berufsgruppen"

    def __str__(self):
        return self.name

    def save_from_post(self, company_id, devices):
        self.company_id = company_id
        self.save()
        self.devices.clear()
        self.devices.add(*devices)


class Employee(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="employees",
    )
    title = models.CharField("Titel", max_length=20, null=True, blank=True)
    first_name = models.CharField("Vorname", max_length=50)
    last_name = models.CharField("Nachname", max_length=50)
    prof_group = models.ManyToManyField(
        ProfessionalGroup,
        verbose_name="Berufsgruppe",
        blank=True,
        related_name="employees",
    )
    is_instructor = models.BooleanField(
        default=False, verbose_name="MPG-Beauftragter"
    )
    instructor_for = models.ManyToManyField(
        Device,
        verbose_name="Einweiser für",
        related_name="instructors",
        blank=True,
    )

    class Meta:
        verbose_name = "Mitarbeiter"
        verbose_name_plural = "Mitarbeiter"
        ordering = ["last_name", "first_name"]

    def __str__(self):
        if self.title:
            return f"{self.title} {self.first_name} {self.last_name}"
        return f"{self.first_name} {self.last_name}"

    def sort_key(self):
        return f"{self.last_name} {self.first_name}"

    def save_from_post(self, company_id, prof_groups):
        self.company_id = company_id
        self.save()
        self.prof_group.clear()
        self.prof_group.add(*prof_groups)

    def lacking_instructions(self):
        devices = Device.objects.filter(professional_groups__employees=self)

    @classmethod
    def get_sorted(cls, company_id):
        """Return Employees for this company in a particular order

        company_id can be an int or a request
        """
        employees = []
        for emp in (
            company_s(company_id, Employee)
            .order_by("last_name", "first_name")
            .prefetch_related("prof_group")
        ):
            pgs = emp.prof_group.all()
            if pgs:
                for pg in pgs:
                    employees.append((pg.name, emp))
            else:
                employees.append((" Keine Gruppe", emp))

        return sorted(employees, key=lambda t: (t[0], t[1].sort_key()))

    def classes(self):
        return " ".join([f"pg_{pg.id}" for pg in self.prof_group.all()])


class Instruction(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="instructions",
    )
    day = models.DateField("Datum", default=date.today)
    instructor = models.ForeignKey(
        Employee,
        verbose_name="Einweiser",
        on_delete=models.PROTECT,
        limit_choices_to={"is_instructor": True},
        related_name="held_instructions",
    )
    instructed = models.ManyToManyField(
        Employee, verbose_name="Eingewiesene", related_name="instructions"
    )
    # TODO: limit to devices for this instructor
    devices = models.ManyToManyField(
        Device, verbose_name="Geräte", related_name="instructions"
    )

    class Meta:
        verbose_name = "Einweisung"
        verbose_name_plural = "Einweisungen"

    def __str__(self):
        return f"Einweisung am {self.day.strftime('%d.%m.%y')} "


class PrimaryInstruction(models.Model):
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="primary_instructions",
    )
    day = models.DateField("Datum", default=date.today)
    instructor = models.CharField("Einweiser", max_length=50)
    device_company = models.CharField("Firma", max_length=50)
    instructed = models.ManyToManyField(
        Employee,
        verbose_name="Eingewiesene",
        related_name="primaryinstructions",
    )
    devices = models.ManyToManyField(
        Device, verbose_name="Geräte", related_name="primaryinstructions"
    )

    class Meta:
        verbose_name = "Ersteinweisung"
        verbose_name_plural = "Ersteinweisungen"

    def __str__(self):
        return f"Ersteinweisung am {self.day.strftime('%d.%m.%y')} "

    def save_instructor_status(self):
        devices = self.devices.all()
        for empl in self.instructed.filter(is_instructor=True):
            empl.instructor_for.add(*devices)


class SiteUser(models.Model):
    user = models.OneToOneField(
        User, on_delete=models.CASCADE, related_name="site_user"
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name="Krankenhaus",
        related_name="site_users",
    )


def company_s(request_or_company_id, model, *args, **kwargs):
    """Return a queryset filtered to the request's company"""
    company_id = (
        request_or_company_id
        if isinstance(request_or_company_id, int)
        else request_or_company_id.session.get("company_id")
    )
    if company_id is None:
        print(f"{company_id=} wurde nicht gefunden")
        raise Http404
    return model.objects.filter(company=company_id, *args, **kwargs)


def save_c(form, company_id):
    model = form.save(commit=False)
    model.company_id = company_id
    if form.instance is None:
        model.id = form.instance.id
    model.save()
    return model


def make_cert(employee):
    """Return a Certification on the instructions for one employee

    Uses 'data/Geraetepass-{company_id}.docx' as template
    Can raise an OSError
    """
    instructions = (
        Instruction.objects.filter(instructed=employee)
        .prefetch_related("devices")
        .select_related("instructor")
    )
    primary_instructions = PrimaryInstruction.objects.filter(
        instructed=employee
    ).prefetch_related("devices")
    dev1 = [
        (dev.vendor, dev.name, instr.day, str(instr.instructor))
        for instr in instructions
        for dev in instr.devices.all()
    ]
    dev2 = [
        (
            dev.vendor,
            dev.name,
            p_instr.day,
            f"{p_instr.instructor}, {p_instr.device_company}",
        )
        for p_instr in primary_instructions
        for dev in p_instr.devices.all()
    ]
    devices = sorted(dev1 + dev2)

    FILENAME = f"Geraetepass-{employee.company_id}.docx"
    TEMPLATE_PATH = settings.BASE_DIR / "data" / FILENAME
    try:
        with open(TEMPLATE_PATH, "rb") as f:
            doc = Document(f)
    except OSError:
        # The caller has to catch
        raise

    # write name and date
    for p in doc.paragraphs:
        if "Geräteeinweisungen für" in p.text:
            p.runs[-1].add_text(str(employee))
        elif "Hann. Münden, den" in p.text:
            p.runs[-1].add_text(date.today().strftime("%d.%m.%Y"))
    # write instructions
    for t in doc.tables:
        first_row = t.rows[0]
        if tuple(c.text.strip() for c in first_row.cells) == (
            "Datum",
            "Einweiser",
            "Geräte",
        ):
            for vendor, name, day, instructor in devices:
                row = t.add_row()
                row.cells[0].text = day.strftime("%d.%m.%Y")
                row.cells[1].text = instructor
                row.cells[2].text = f"{vendor}: {name}"
            if not devices:
                row = t.add_row()
                row.cells[2].text = f"Noch keine Einweisungen"
            break

    filepath = settings.BASE_DIR / "data" / f"Geraetepass {employee}.docx"
    doc.save(filepath)

    return filepath
